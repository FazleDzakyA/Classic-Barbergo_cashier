import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import re

def create_proposal_docx(md_path, docx_path):
    doc = Document()

    # Set page margins (Standard A4 / 1 inch margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    # Styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def process_table(lines):
        if not lines:
            return
        rows = [line.strip().strip('|').split('|') for line in lines if not re.match(r'^\s*\|?\s*:?-+:?\s*\|', line)]
        if not rows:
            return

        cols_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < cols_count:
                    cell = table.cell(i, j)
                    cell.text = cell_text.strip()
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)

                    if i == 0:  # Header row
                        set_cell_background(cell, "1E293B")
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    else:
                        if i % 2 == 1:
                            set_cell_background(cell, "F8FAFC")
                        else:
                            set_cell_background(cell, "FFFFFF")
                        for run in p.runs:
                            run.font.size = Pt(9.5)
        doc.add_paragraph()  # Spacing

    def process_code_block(code_lines):
        text = "".join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    for line in lines:
        raw_line = line.rstrip('\r\n')
        stripped = raw_line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                process_code_block(code_block_lines)
                code_block_lines = []
                in_code_block = False
            else:
                if in_table:
                    process_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(raw_line + '\n')
            continue

        # Table handling
        if '|' in stripped and stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        else:
            if in_table:
                process_table(table_lines)
                table_lines = []
                in_table = False

        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=1)
            h.style.font.name = 'Arial'
            h.style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            h.style.font.name = 'Arial'
            h.style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            h.style.font.name = 'Arial'
            h.style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        elif stripped.startswith('#### '):
            h = doc.add_heading(stripped[5:], level=4)
            h.style.font.name = 'Arial'
            h.style.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Bold inline parsing
            parts = re.split(r'(\*\*.*?\*\*)', stripped[2:])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            content = re.sub(r'^\d+\.\s', '', stripped)
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
        elif stripped.startswith('---'):
            doc.add_paragraph('—' * 40)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                else:
                    p.add_run(part)

    if in_table:
        process_table(table_lines)

    try:
        doc.save(docx_path)
        print(f"Successfully created docx at {docx_path}")
    except PermissionError:
        fallback_path = docx_path.replace('.docx', '_REVISI_RAB.docx')
        doc.save(fallback_path)
        print(f"File locked by Word. Saved to {fallback_path}")

if __name__ == '__main__':
    create_proposal_docx('PROPOSAL_BARBERFLOW.md', 'PROPOSAL_BARBERFLOW.docx')
